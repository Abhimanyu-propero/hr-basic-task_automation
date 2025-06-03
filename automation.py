from flask import Flask, render_template, request, send_file
from docx import Document
from jinja2 import Template
import os
import smtplib
from email.message import EmailMessage

app = Flask(__name__)

# Fill Word template with values
def fill_docx(template_path, output_path, context):
    doc = Document(template_path)

    # Replace in paragraphs
    for paragraph in doc.paragraphs:
        full_text = paragraph.text
        for key, value in context.items():
            full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
        # Rebuild paragraph with replaced text
        if paragraph.text != full_text:
            for run in paragraph.runs:
                run.text = ''
            paragraph.runs[0].text = full_text if paragraph.runs else paragraph.add_run(full_text)

    # Replace in tables too (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    full_text = paragraph.text
                    for key, value in context.items():
                        full_text = full_text.replace(f"{{{{{key}}}}}", str(value))
                    if paragraph.text != full_text:
                        for run in paragraph.runs:
                            run.text = ''
                        paragraph.runs[0].text = full_text if paragraph.runs else paragraph.add_run(full_text)

    doc.save(output_path)
# Send email with attachment
def send_email_with_attachment(to_email, subject, body, attachment_path):
    msg = EmailMessage()
    msg["Subject"] = subject
    msg["From"] = "abhimanyusingh877@gmail.com"
    msg["To"] = to_email
    msg.set_content(body)

    with open(attachment_path, "rb") as f:
        file_data = f.read()
        file_name = os.path.basename(attachment_path)
        msg.add_attachment(file_data, maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=file_name)

    with smtplib.SMTP_SSL("smtp.gmail.com", 465) as smtp:
        smtp.login("abhimanyusingh877@gmail.com", "iwazdduchqzmzwqf")
        smtp.send_message(msg)

@app.route("/")
def home():
    return render_template("main.html")

@app.route("/appointment_letter.html", methods=["GET", "POST"])
def appointment():
    if request.method == "POST":
        data = request.form.to_dict()
        data["name"] = f"{data['first_name']} {data['last_name']}"
        name = data["name"].replace(" ", "_")  # Sanitize filename
        output_path = f"{data['first_name']}_appointment_letter.docx"
        fill_docx("Appointment_letter.docx", output_path, data)
        send_email_with_attachment(data["email"], "Your Appointment Letter", "Please find the attached letter.", output_path)
        return f"Appointment letter sent to {data['email']}!"
    return render_template("appointment_letter.html")

@app.route("/experience_letter.html", methods=["GET", "POST"])
def experience():
    if request.method == "POST":
        data = request.form.to_dict()
        name = data["name"].replace(" ", "_")  # Sanitize filename
        output_path = f"{name}_experience_letter.docx"
        fill_docx("Experience_letter.docx", output_path, data)
        send_email_with_attachment(data["email"], "Your Experience Letter", "Please find the attached letter.", output_path)
        return f"Experience letter sent to {data['email']}!"
    return render_template("experience_letter.html")

@app.route("/internship_letter.html", methods=["GET", "POST"])
def internship():
    if request.method == "POST":
        data = request.form.to_dict()
        name = data["name"].replace(" ", "_")  # Sanitize filename
        output_path = f"{name}_internship_letter.docx"
        fill_docx("Internship_letter.docx", output_path, data)
        send_email_with_attachment(data["email"], "Your Internship Letter", "Please find the attached letter.", output_path)
        return f"Internship letter sent to {data['email']}!"
    return render_template("internship_letter.html")

if __name__ == "__main__":
    app.run(debug=True)
